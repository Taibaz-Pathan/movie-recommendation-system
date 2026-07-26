from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, name='Calibri', size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(doc, text='', size=11, bold=False, space_before=0,
                  space_after=8, color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.alignment = alignment
    p.paragraph_format.line_spacing = Pt(18)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    if level == 1:
        p = add_paragraph(doc, text, size=16, bold=True,
                          space_before=20, space_after=6, color=(17, 17, 17))
    else:
        p = add_paragraph(doc, text, size=12, bold=True,
                          space_before=18, space_after=4, color=(30, 30, 30))
        # bottom border on h2
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), 'CCCCCC')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(text)
    set_font(run, size=size)
    return p


def add_divider(doc):
    p = add_paragraph(doc, space_before=6, space_after=6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'DDDDDD')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_inline(doc, parts, size=11, space_before=0, space_after=8):
    """Add a paragraph with mixed bold/normal runs. parts = [(text, bold), ...]"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(18)
    for text, bold in parts:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            set_font(run, size=10, bold=True)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'F0F0F0')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = val
            for run in row_cells[c_idx].paragraphs[0].runs:
                bold = val.startswith('**') or (r_idx == len(rows) - 1)
                set_font(run, size=10, bold=bold)

    doc.add_paragraph()


# ── Build document ──────────────────────────────────────────────────────────

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1.1)
    section.bottom_margin = Inches(1.1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── Header info ─────────────────────────────────────────────────────────────

add_paragraph(doc, 'Project Progress Report — Weeks 1 to 4',
              size=18, bold=True, space_after=10, color=(17, 17, 17))

meta = [
    ('Project Title: ', 'Movie Recommendation System Using Collaborative Filtering'),
    ('Student: ', 'Taibaz Pathan'),
    ('Student ID: ', '284085'),
    ('University: ', 'Frankfurt University of Applied Sciences'),
    ('Supervisor: ', 'Prof. Dr. Andreas Pech'),
    ('Report Date: ', '02 June 2026'),
    ('Project Start: ', '07 May 2026'),
]
for label, value in meta:
    add_inline(doc, [(label, True), (value, False)], space_after=3)

add_divider(doc)

# ── Overview ────────────────────────────────────────────────────────────────

add_heading(doc, 'Overview', level=2)
add_paragraph(doc,
    'This report covers the first four weeks of work on my movie recommendation '
    'system project. The aim is to build and compare User-Based and Item-Based '
    'Collaborative Filtering models on the MovieLens dataset, and ultimately '
    'evaluate them using standard accuracy metrics. These opening weeks focused '
    'on getting the foundations right: setting up a clean project structure, '
    'understanding the data properly, preprocessing it for modelling, and '
    'implementing the core similarity functions that the CF models will rely on.',
    space_after=10)

# ── Week 1 ──────────────────────────────────────────────────────────────────

add_heading(doc, 'Week 1 — Project Setup and Environment', level=2)
add_paragraph(doc,
    'The first week was spent building the project scaffold and getting the '
    'development environment into a consistent, reproducible state. I set up a '
    'virtual environment and pinned all dependencies to specific versions to '
    'avoid compatibility issues later. The dataset used throughout is the '
    'MovieLens ml-latest-small release from GroupLens Research, which contains '
    '100,836 ratings from 610 users across 9,742 movies.',
    space_after=6)
add_paragraph(doc,
    'The main deliverables from this week were:', space_after=4)
add_bullet(doc, 'A structured folder layout following Python packaging conventions, '
           'with separate directories for data, source code, notebooks, tests, and reports.')
add_bullet(doc, 'A centralised configuration file (configs/config.yaml) for dataset '
           'paths, split parameters, and logging settings.')
add_bullet(doc, 'A data loading module (src/data/loader.py) that reads the MovieLens '
           'CSV files and automatically converts the Unix timestamp column to a '
           'human-readable datetime.')
add_bullet(doc, 'An evaluation metrics module (src/evaluation/metrics.py) containing '
           'RMSE and MAE functions with full input validation.')
add_bullet(doc, 'Skeleton classes for the two CF models that will be fully implemented '
           'in Weeks 5 and 6.')
add_bullet(doc, 'A test suite with 17 passing unit tests covering the loader and '
           'metrics modules.')
add_paragraph(doc, 'The Git repository was initialised on branch main with the '
              'first commit at the end of the week.', space_before=4, space_after=10)

# ── Week 2 ──────────────────────────────────────────────────────────────────

add_heading(doc, 'Week 2 — Exploratory Data Analysis', level=2)
add_paragraph(doc,
    'Before building any models, I wanted to properly understand what the data '
    'looks like. A Jupyter notebook (notebooks/eda.ipynb) was produced covering '
    'eight analysis areas, with seven figures saved to reports/figures/. '
    'The most important findings are summarised below.',
    space_after=8)

add_inline(doc, [('Rating distribution.  ', True),
    ('The mean rating is 3.50, which is above the theoretical midpoint of 2.75. '
     'Users show a clear positive bias — they tend to rate films they already '
     'chose to watch, so low ratings are underrepresented. Whole-star values '
     '(1, 2, 3, 4, 5) appear far more often than half-star values. '
     'This bias will need to be corrected through mean-centring in the CF models.', False)],
    space_after=6)

add_inline(doc, [('User activity.  ', True),
    ('The distribution of ratings per user is heavily skewed. One user has over '
     '2,000 ratings while the median user has only 68. This matters for UBCF: '
     'users with few ratings will produce unreliable Pearson correlation scores '
     'because there are too few co-rated movies to compute a meaningful similarity.', False)],
    space_after=6)

add_inline(doc, [('Movie popularity.  ', True),
    ('A handful of blockbuster films dominate the dataset. The median movie has '
     'only 3 ratings in total, which is too few for any similarity computation '
     'to be reliable. I will filter out movies with fewer than 20 ratings '
     'before training the models.', False)],
    space_after=6)

add_inline(doc, [('Matrix sparsity.  ', True),
    ('Only 1.7% of the full user-item matrix is filled — the remaining 98.3% '
     'is empty. This means most pairs of users share very few co-rated movies, '
     'which makes Pearson correlation unstable without a minimum co-rating '
     'threshold (min_support).', False)],
    space_after=6)

add_inline(doc, [('Temporal patterns.  ', True),
    ('Ratings span 1996 to 2018. Activity peaked around 2000 and the average '
     'rating per year has remained close to 3.5 throughout, suggesting no '
     'significant drift in user behaviour over time.', False)],
    space_after=10)

# ── Week 3 ──────────────────────────────────────────────────────────────────

add_heading(doc, 'Week 3 — Data Preprocessing', level=2)
add_paragraph(doc,
    'With a clear picture of the data, I built a preprocessing module '
    '(src/data/preprocessor.py) to clean the dataset and prepare it for '
    'model training. The module has four functions.',
    space_after=8)

add_inline(doc, [('Filtering.  ', True),
    ('filter_ratings() removes users and movies that fall below a minimum '
     'rating count. The filtering runs iteratively until convergence — removing '
     'sparse movies can make some users drop below the threshold, so a single '
     'pass is not enough. With thresholds of 20 ratings per user and per movie, '
     'the dataset shrinks from 100,836 to 67,020 ratings across 566 users '
     'and 1,286 movies.', False)],
    space_after=6)

add_inline(doc, [('User-item matrix.  ', True),
    ('build_user_item_matrix() constructs a pivot table with users as rows '
     'and movies as columns. Unrated entries are NaN. After filtering, the '
     'matrix is 566 x 1,286 with 92.6% sparsity — lower than the raw dataset '
     'because rare movies have been removed.', False)],
    space_after=6)

add_inline(doc, [('Train/test split.  ', True),
    ('train_test_split_per_user() holds out 20% of each user\'s ratings '
     'as the test set. Splitting per user (rather than globally) guarantees '
     'that every user appears in both splits, which is essential for CF '
     'evaluation. The result is 53,614 training ratings and 13,406 test '
     'ratings, with the random seed fixed at 42 for reproducibility.', False)],
    space_after=6)

add_inline(doc, [('Saving splits.  ', True),
    ('save_splits() writes train.csv and test.csv to data/processed/, '
     'so the same split can be reused across all model experiments without '
     'recomputing it each time.', False)],
    space_after=6)

add_paragraph(doc, '18 unit tests were written for this module, all passing.',
              space_after=10)

# ── Week 4 ──────────────────────────────────────────────────────────────────

add_heading(doc, 'Week 4 — Similarity Metrics from Scratch', level=2)
add_paragraph(doc,
    'This week I implemented the core similarity functions that the CF models '
    'will use, entirely from scratch in NumPy — no pandas .corr() or '
    'scikit-learn. The module lives at src/utils/similarity.py and provides '
    'both vector-level functions (useful for teaching and single-pair '
    'calculations) and vectorised matrix-level functions for computing '
    'full similarity matrices efficiently.',
    space_after=8)

add_inline(doc, [('cosine_similarity(u, v).  ', True),
    ('Computes cosine similarity between two rating vectors. Only positions '
     'where both vectors are non-NaN are used, so missing ratings are excluded '
     'rather than treated as zeros.', False)],
    space_after=6)

add_inline(doc, [('pearson_similarity(u, v, min_support).  ', True),
    ('Computes the exact Pearson correlation, mean-centring each vector over '
     'the co-rated items only — not over each user\'s full rating history. '
     'This matches the formulation from the original CF literature '
     '(Resnick et al., 1994). If the number of shared ratings is below '
     'min_support, the function returns 0.0 rather than an unreliable score.', False)],
    space_after=6)

add_inline(doc, [('cosine_similarity_matrix(matrix).  ', True),
    ('Vectorised computation of the full pairwise cosine similarity matrix. '
     'NaN entries are replaced with 0 before L2 normalisation, then the '
     'similarity is computed as a single matrix multiplication.', False)],
    space_after=6)

add_inline(doc, [('pearson_similarity_matrix(matrix, min_support).  ', True),
    ('Vectorised Pearson similarity matrix. Each row is mean-centred using '
     'its own overall mean rating, and pairs with fewer co-rated items than '
     'min_support are zeroed out. This runs efficiently on the full '
     '566 x 1,286 matrix without nested Python loops.', False)],
    space_after=6)

add_paragraph(doc, '22 unit tests were written for this module, all passing.',
              space_after=10)

add_divider(doc)

# ── Test summary table ───────────────────────────────────────────────────────

add_heading(doc, 'Test Suite Summary', level=2)
add_table(doc,
    headers=['Module', 'Tests', 'Status'],
    rows=[
        ['src/data/loader.py',        '7',    'Passing'],
        ['src/evaluation/metrics.py', '10',   'Passing'],
        ['src/data/preprocessor.py',  '18',   'Passing'],
        ['src/utils/similarity.py',   '22',   'Passing'],
        ['Total',                     '57',   'All passing'],
    ])

# ── Repository structure ─────────────────────────────────────────────────────

add_heading(doc, 'Repository Structure', level=2)
add_paragraph(doc,
    'The project follows a standard Python package layout. All source code '
    'is under src/, tests mirror that structure under tests/, and generated '
    'outputs (figures, processed data) are kept separate from the code.',
    space_after=8)

structure_lines = [
    'movie-recsys/',
    '├── configs/config.yaml            centralised configuration',
    '├── data/processed/                train.csv, test.csv',
    '├── notebooks/eda.ipynb            Week 2 EDA',
    '├── reports/figures/               7 EDA figures (PNG)',
    '├── src/',
    '│   ├── data/loader.py             dataset loading',
    '│   ├── data/preprocessor.py       filtering, matrix, split',
    '│   ├── models/ubcf.py             User-Based CF (Week 5)',
    '│   ├── models/ibcf.py             Item-Based CF (Week 6)',
    '│   ├── evaluation/metrics.py      RMSE, MAE',
    '│   └── utils/similarity.py        cosine and Pearson from scratch',
    '└── tests/                         57 unit tests',
]
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(10)
p.paragraph_format.line_spacing = Pt(16)
for line in structure_lines:
    run = p.add_run(line + '\n')
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(50, 50, 50)

# ── Plan for weeks 5 and 6 ──────────────────────────────────────────────────

add_heading(doc, 'Plan for Weeks 5 and 6', level=2)
add_inline(doc, [('Week 5.  ', True),
    ('Complete the User-Based CF model by wiring the custom similarity module '
     'into the UBCF class. The model will be trained on the training split and '
     'evaluated on the test split using RMSE and MAE.', False)],
    space_after=6)
add_inline(doc, [('Week 6.  ', True),
    ('Do the same for Item-Based CF and run an initial side-by-side comparison '
     'of UBCF versus IBCF accuracy.', False)],
    space_after=10)

add_divider(doc)
add_paragraph(doc,
    'All code is version-controlled in a local Git repository. '
    'Four commits have been made, one per week milestone.',
    size=10, color=(100, 100, 100), space_after=0)

doc.save('reports/progress_week1_4.docx')
print('Done — reports/progress_week1_4.docx')
